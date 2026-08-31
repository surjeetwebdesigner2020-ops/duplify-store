from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pathlib import Path
import re

SRC = Path(r"C:\Users\HP\Desktop\store-duplicate\Store-Duplicate-Troubleshooting-Guide-Hindi.docx")
OUT = Path(r"C:\Users\HP\Desktop\store-duplicate\Store-Duplicate-Troubleshooting-Guide-Hinglish.docx")

M = {
"मुख्य निष्कर्ष: App code खराब नहीं था। गलत Shopify app credentials, पुराना Railway domain और local DNS refusal—इन तीनों के कारण embedded app blank दिख रहा था।": "Main conclusion: App code kharab nahi tha. Galat Shopify app credentials, purana Railway domain aur local DNS refusal - in teen problems ki wajah se embedded app blank dikh raha tha.",
"इस document का उपयोग कब करें": "Is document ko kab use karein",
"Shopify Admin में app blank page या broken-file icon दिखाए।": "Shopify Admin mein app blank page ya broken-file icon dikhaye.",
"‘server IP address could not be found’ या DNS error आए।": "'server IP address could not be found' ya DNS error aaye.",
"Railway service Online हो, लेकिन app iframe न खुले।": "Railway service Online ho, lekin app iframe na khule.",
"OAuth/login loop, invalid API key या callback mismatch आए।": "OAuth/login loop, invalid API key ya callback mismatch aaye.",
"Redis/BullMQ या PostgreSQL connection error आए।": "Redis/BullMQ ya PostgreSQL connection error aaye.",
"सुरक्षा: इस file में कोई password, API secret, database password या Redis password नहीं है। Secrets केवल Railway Variables में रखें।": "Security: Is file mein koi password, API secret, database password ya Redis password nahi hai. Secrets sirf Railway Variables mein rakhein.",
"1. असली issue क्या था?": "1. Asli issue kya tha?",
"Permanent fix जो किया गया": "Permanent fix jo kiya gaya",
"Project config को Shopify के ‘store duplicate’ app से link किया गया।": "Project config ko Shopify ke 'store duplicate' app se link kiya gaya.",
"Broken Railway domain हटाकर नया production domain बनाया गया।": "Broken Railway domain hata kar naya production domain banaya gaya.",
"Shopify App Home URL और OAuth callback नए domain पर release किए गए।": "Shopify App Home URL aur OAuth callback naye domain par release kiye gaye.",
"Railway में सही SHOPIFY_APP_URL, SHOPIFY_API_KEY और SHOPIFY_API_SECRET set किए गए।": "Railway mein sahi SHOPIFY_APP_URL, SHOPIFY_API_KEY aur SHOPIFY_API_SECRET set kiye gaye.",
"Wi-Fi DNS को 1.1.1.1 और 8.8.8.8 पर बदला गया और DNS cache साफ किया गया।": "Wi-Fi DNS ko 1.1.1.1 aur 8.8.8.8 par badla gaya aur DNS cache clear kiya gaya.",
"GitHub main branch और Railway deployment verify किए गए।": "GitHub main branch aur Railway deployment verify kiye gaye.",
"2. सबसे तेज recovery checklist": "2. Sabse fast recovery checklist",
"पहले यह करें: Shopify page पर Ctrl + Shift + R से hard refresh करें। फिर नीचे दिए checks क्रम से चलाएँ।": "Sabse pehle: Shopify page par Ctrl + Shift + R se hard refresh karein. Phir neeche diye checks ko sequence mein chalayein.",
"Public URL खोलें — नया Railway URL browser में खोलें। 200 response या app page आना चाहिए।": "Public URL kholein - Naya Railway URL browser mein kholein. 200 response ya app page aana chahiye.",
"Railway deployment देखें — App service का latest deployment SUCCESS और replica RUNNING होना चाहिए।": "Railway deployment dekhein - App service ka latest deployment SUCCESS aur replica RUNNING hona chahiye.",
"Domain status देखें — Domain ACTIVE हो और target port 3000 हो।": "Domain status dekhein - Domain ACTIVE ho aur target port 3000 ho.",
"Shopify URLs मिलाएँ — App Home URL और redirect URL उसी exact HTTPS hostname पर हों; callback path /auth/callback हो।": "Shopify URLs match karein - App Home URL aur redirect URL same exact HTTPS hostname par hon; callback path /auth/callback ho.",
"Credentials मिलाएँ — Railway की SHOPIFY_API_KEY और SHOPIFY_API_SECRET उसी Shopify app से हों जिसका नाम ‘store duplicate’ है।": "Credentials match karein - Railway ki SHOPIFY_API_KEY aur SHOPIFY_API_SECRET usi Shopify app ke hon jiska naam 'store duplicate' hai.",
"Database/Redis देखें — PostgreSQL और Redis services Online हों; internal URLs Railway Variables में हों।": "Database/Redis dekhein - PostgreSQL aur Redis services Online hon; internal URLs Railway Variables mein hon.",
"Logs पढ़ें — Latest runtime logs में server 0.0.0.0:3000 पर listen कर रहा हो और migration error न हो।": "Logs padhein - Latest runtime logs mein server 0.0.0.0:3000 par listen kar raha ho aur migration error na ho.",
"3. Error के हिसाब से solution": "3. Error ke hisaab se solution",
"पहचान: Railway URL open नहीं होता; nslookup/Resolve-DnsName fail या REFUSED देता है।": "Pehchan: Railway URL open nahi hota; nslookup/Resolve-DnsName fail ya REFUSED deta hai.",
"Fix: Public DNS 1.1.1.1 से test करें। केवल local DNS fail हो तो adapter DNS 1.1.1.1 और 8.8.8.8 करें, cache clear करें। Domain globally fail हो तो Railway domain status check करें।": "Fix: Public DNS 1.1.1.1 se test karein. Sirf local DNS fail ho to adapter DNS 1.1.1.1 aur 8.8.8.8 karein, phir cache clear karein. Domain globally fail ho to Railway domain status check karein.",
"पहचान: App title दिखता है, content area blank रहता है।": "Pehchan: App title dikhta hai, lekin content area blank rehta hai.",
"Fix: App Home URL, callback URL, DNS और HTTP response check करें। Browser hard refresh करें। Railway logs में request आनी चाहिए।": "Fix: App Home URL, callback URL, DNS aur HTTP response check karein. Browser hard refresh karein. Railway logs mein request aani chahiye.",
"पहचान: Login loop, unauthorized, invalid client या callback mismatch।": "Pehchan: Login loop, unauthorized, invalid client ya callback mismatch.",
"Fix: API key + secret को उसी Shopify app से लें। Shopify config और Railway variables को एक ही app identity पर रखें। Secret कभी Git में commit न करें।": "Fix: API key + secret usi Shopify app se lein. Shopify config aur Railway variables ko ek hi app identity par rakhein. Secret kabhi Git mein commit na karein.",
"पहचान: Prisma migration या database connection fail।": "Pehchan: Prisma migration ya database connection fail.",
"Fix: DATABASE_URL Railway PostgreSQL internal hostname का हो। Postgres service Online हो। Existing volume को delete न करें।": "Fix: DATABASE_URL Railway PostgreSQL internal hostname ka ho. Postgres service Online ho. Existing volume ko delete na karein.",
"पहचान: Queue शुरू नहीं होती या ECONNREFUSED 6379।": "Pehchan: Queue start nahi hoti ya ECONNREFUSED 6379 aata hai.",
"Fix: REDIS_URL Railway Redis internal hostname/password का हो। Redis service और volume Online हों।": "Fix: REDIS_URL Railway Redis internal hostname/password ka ho. Redis service aur volume Online hon.",
"4. Config values जो हमेशा match होने चाहिए": "4. Config values jo hamesha match hone chahiye",
"क्या कभी delete नहीं करना": "Kin cheezon ko kabhi delete nahi karna",
"PostgreSQL service या postgres-volume, जब तक verified backup और explicit migration plan न हो।": "PostgreSQL service ya postgres-volume, jab tak verified backup aur clear migration plan na ho.",
"Redis volume बिना यह समझे कि pending jobs/idempotency state पर क्या प्रभाव पड़ेगा।": "Redis volume ko tab tak delete na karein jab tak pending jobs aur idempotency state ka impact clear na ho.",
".env या secrets को GitHub पर push न करें।": ".env ya secrets ko GitHub par push na karein.",
"Working domain delete न करें; पहले replacement/rollback plan रखें।": "Working domain delete na karein; pehle replacement/rollback plan rakhein.",
"Secret leak response: यदि कोई Shopify secret, database password या Redis password chat/screenshot में share हो जाए, उसे exposed मानें और provider dashboard से rotate करें।": "Secret leak response: Agar koi Shopify secret, database password ya Redis password chat/screenshot mein share ho jaye, use exposed maanein aur provider dashboard se rotate karein.",
"5. Final health check": "5. Final health check",
"☐ Shopify Admin में store duplicate app dashboard खुलता है।": "☐ Shopify Admin mein store duplicate app dashboard khulta hai.",
"☐ Railway app deployment SUCCESS है और replica RUNNING है।": "☐ Railway app deployment SUCCESS hai aur replica RUNNING hai.",
"☐ Public URL HTTP 200 देता है।": "☐ Public URL HTTP 200 deta hai.",
"☐ PostgreSQL migration: No pending migrations / no errors।": "☐ PostgreSQL migration: No pending migrations / no errors.",
"☐ Redis connection error नहीं है।": "☐ Redis connection error nahi hai.",
"☐ Latest logs में repeated crash, 401, 500 या DNS errors नहीं हैं।": "☐ Latest logs mein repeated crash, 401, 500 ya DNS errors nahi hain.",
"Project पहले ‘Duplify Store’ client ID से linked था; target ‘store duplicate’ app था।": "Project pehle 'Duplify Store' client ID se linked tha; target 'store duplicate' app tha.",
"गलत app identity / OAuth mismatch": "Galat app identity / OAuth mismatch",
"पुराना public hostname DNS पर resolve नहीं हो रहा था।": "Purana public hostname DNS par resolve nahi ho raha tha.",
"Wi-Fi DNS server ने Railway domain के लिए REFUSED दिया।": "Wi-Fi DNS server ne Railway domain ke liye REFUSED diya.",
"Browser domain नहीं खोल पाया": "Browser domain nahi khol paya",
"SHOPIFY_API_KEY / SECRET पुराने app के थे।": "SHOPIFY_API_KEY / SECRET purane app ke the.",
"Authentication fail या login loop": "Authentication fail ya login loop",
"PostgreSQL और Redis healthy थे।": "PostgreSQL aur Redis healthy the.",
"इनको delete/recreate करने की जरूरत नहीं थी": "Inko delete/recreate karne ki zarurat nahi thi",
"3000; server 0.0.0.0 पर listen करे।": "3000; server 0.0.0.0 par listen kare.",
"Current Railway HTTPS public domain; trailing slash के बिना।": "Current Railway HTTPS public domain; trailing slash ke bina.",
"SHOPIFY_APP_URL के exact बराबर।": "SHOPIFY_APP_URL ke exact barabar.",
"‘store duplicate’ Shopify app की Client ID।": "'store duplicate' Shopify app ki Client ID.",
"उसी ‘store duplicate’ app का secret; chat/Git में कभी न रखें।": "Ussi 'store duplicate' app ka secret; chat/Git mein kabhi na rakhein.",
"Railway Postgres internal URL; app service से accessible।": "Railway Postgres internal URL; app service se accessible.",
"Railway Redis internal URL; app service से accessible।": "Railway Redis internal URL; app service se accessible.",
}

def all_paragraphs(document):
    for p in document.paragraphs:
        yield p
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

doc = Document(SRC)
for p in all_paragraphs(doc):
    old = p.text
    if old in M:
        p.text = M[old]
        for run in p.runs:
            run.font.name = "Calibri"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
            run.font.size = Pt(11)

remaining = []
for p in all_paragraphs(doc):
    if re.search(r"[\u0900-\u097F]", p.text):
        remaining.append(p.text)
if remaining:
    raise RuntimeError("Unconverted Devanagari: " + " | ".join(remaining))

doc.core_properties.title = "Store Duplicate Troubleshooting & Recovery Guide - Hinglish"
doc.save(OUT)
print(OUT)
