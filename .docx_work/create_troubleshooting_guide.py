from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\HP\Desktop\store-duplicate\Store-Duplicate-Troubleshooting-Guide-Hindi.docx")

BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
RED = "9B1C1C"
GREEN = "2E6B3F"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

def set_font(run, size=11, bold=None, color=None, italic=None):
    run.font.name = "Nirmala UI"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Nirmala UI")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Nirmala UI")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Nirmala UI")
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Nirmala UI"; normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Nirmala UI")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Nirmala UI")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala UI")
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK, 10, 5),
]:
    st = styles[name]
    st.font.name = "Nirmala UI"; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn("w:ascii"), "Nirmala UI")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Nirmala UI")
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala UI")
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ["List Bullet", "List Number"]:
    st = styles[list_name]
    st.font.name = "Nirmala UI"; st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.375)
    st.paragraph_format.first_line_indent = Inches(-0.188)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
callout.base_style = normal
callout.paragraph_format.left_indent = Inches(0.18)
callout.paragraph_format.right_indent = Inches(0.18)
callout.paragraph_format.space_before = Pt(6)
callout.paragraph_format.space_after = Pt(8)

def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); pPr.append(shd)

def add_callout(label, text, fill=LIGHT, color=DARK):
    p = doc.add_paragraph(style="Callout"); shade_paragraph(p, fill)
    r = p.add_run(label + " "); set_font(r, 11, True, color)
    r = p.add_run(text); set_font(r, 11)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet"); set_font(p.add_run(text), 11); return p

def add_step(title, detail):
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run(title + " — "), 11, True, DARK)
    set_font(p.add_run(detail), 11)

def set_cell(cell, text, bold=False, fill=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(text), 10.2, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top",80),("bottom",80),("start",120),("end",120)]:
        node=OxmlElement("w:"+side); node.set(qn("w:w"),str(val)); node.set(qn("w:type"),"dxa"); tcMar.append(node)
    cell._tc.get_or_add_tcPr().append(tcMar)

def set_table_widths(table, widths_dxa):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW") or OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths_dxa))); tblW.set(qn("w:type"), "dxa")
    if tblW.getparent() is None: tblPr.append(tblW)
    tblInd = OxmlElement("w:tblInd"); tblInd.set(qn("w:w"),"120"); tblInd.set(qn("w:type"),"dxa"); tblPr.append(tblInd)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths_dxa:
        col=OxmlElement("w:gridCol"); col.set(qn("w:w"),str(w)); grid.append(col)
    for row in table.rows:
        for cell,w in zip(row.cells,widths_dxa):
            tcW=cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:w"),str(w)); tcW.set(qn("w:type"),"dxa")

# Header/footer
hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run("STORE DUPLICATE | RECOVERY GUIDE"), 8.5, True, "6B7280")
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(fp.add_run("Duplify / Store Duplicate production support • बिना secrets"), 8.5, color="6B7280")

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(72); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("STORE DUPLICATE"), 12, True, BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
set_font(p.add_run("Troubleshooting & Recovery Guide"), 28, True, DARK)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(28)
set_font(p.add_run("Shopify + Railway + PostgreSQL + Redis"), 14, color="4B5563")
add_callout("मुख्य निष्कर्ष:", "App code खराब नहीं था। गलत Shopify app credentials, पुराना Railway domain और local DNS refusal—इन तीनों के कारण embedded app blank दिख रहा था।", "EAF3FF")

doc.add_paragraph("इस document का उपयोग कब करें", style="Heading 1")
for x in [
    "Shopify Admin में app blank page या broken-file icon दिखाए।",
    "‘server IP address could not be found’ या DNS error आए।",
    "Railway service Online हो, लेकिन app iframe न खुले।",
    "OAuth/login loop, invalid API key या callback mismatch आए।",
    "Redis/BullMQ या PostgreSQL connection error आए।",
]: add_bullet(x)
add_callout("सुरक्षा:", "इस file में कोई password, API secret, database password या Redis password नहीं है। Secrets केवल Railway Variables में रखें।", "FFF4E5", RED)

doc.add_page_break()
doc.add_paragraph("1. असली issue क्या था?", style="Heading 1")

table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
set_table_widths(table, [2100, 4380, 2880])
for c,t in zip(table.rows[0].cells,["Layer","Problem","Effect"]): set_cell(c,t,True,LIGHT,DARK)
rows = [
    ("Shopify config", "Project पहले ‘Duplify Store’ client ID से linked था; target ‘store duplicate’ app था।", "गलत app identity / OAuth mismatch"),
    ("Railway domain", "पुराना public hostname DNS पर resolve नहीं हो रहा था।", "Shopify iframe blank / broken icon"),
    ("Windows DNS", "Wi-Fi DNS server ने Railway domain के लिए REFUSED दिया।", "Browser domain नहीं खोल पाया"),
    ("Railway variables", "SHOPIFY_API_KEY / SECRET पुराने app के थे।", "Authentication fail या login loop"),
    ("Data services", "PostgreSQL और Redis healthy थे।", "इनको delete/recreate करने की जरूरत नहीं थी"),
]
for a,b,c in rows:
    cells=table.add_row().cells
    for cell,text in zip(cells,[a,b,c]): set_cell(cell,text)

doc.add_paragraph("Permanent fix जो किया गया", style="Heading 2")
for x in [
    "Project config को Shopify के ‘store duplicate’ app से link किया गया।",
    "Broken Railway domain हटाकर नया production domain बनाया गया।",
    "Shopify App Home URL और OAuth callback नए domain पर release किए गए।",
    "Railway में सही SHOPIFY_APP_URL, SHOPIFY_API_KEY और SHOPIFY_API_SECRET set किए गए।",
    "Wi-Fi DNS को 1.1.1.1 और 8.8.8.8 पर बदला गया और DNS cache साफ किया गया।",
    "GitHub main branch और Railway deployment verify किए गए।",
]: add_bullet(x)

doc.add_page_break()
doc.add_paragraph("2. सबसे तेज recovery checklist", style="Heading 1")
add_callout("पहले यह करें:", "Shopify page पर Ctrl + Shift + R से hard refresh करें। फिर नीचे दिए checks क्रम से चलाएँ।")

add_step("Public URL खोलें", "नया Railway URL browser में खोलें। 200 response या app page आना चाहिए।")
add_step("Railway deployment देखें", "App service का latest deployment SUCCESS और replica RUNNING होना चाहिए।")
add_step("Domain status देखें", "Domain ACTIVE हो और target port 3000 हो।")
add_step("Shopify URLs मिलाएँ", "App Home URL और redirect URL उसी exact HTTPS hostname पर हों; callback path /auth/callback हो।")
add_step("Credentials मिलाएँ", "Railway की SHOPIFY_API_KEY और SHOPIFY_API_SECRET उसी Shopify app से हों जिसका नाम ‘store duplicate’ है।")
add_step("Database/Redis देखें", "PostgreSQL और Redis services Online हों; internal URLs Railway Variables में हों।")
add_step("Logs पढ़ें", "Latest runtime logs में server 0.0.0.0:3000 पर listen कर रहा हो और migration error न हो।")

doc.add_paragraph("Safe commands", style="Heading 2")
commands = [
    "railway status --json",
    "railway service list --json",
    "railway deployment list --service duplify-store --limit 5 --json",
    "railway domain list --service duplify-store --json",
    "railway logs --service duplify-store --latest --lines 200 --json",
    "shopify app config validate --json",
]
for cmd in commands:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(cmd); set_font(r, 9.5, color="111827"); r.font.name="Consolas"; r._element.rPr.rFonts.set(qn("w:ascii"),"Consolas"); r._element.rPr.rFonts.set(qn("w:hAnsi"),"Consolas")

doc.add_page_break()
doc.add_paragraph("3. Error के हिसाब से solution", style="Heading 1")

for title, symptoms, fix in [
    ("DNS / server IP not found", "Railway URL open नहीं होता; nslookup/Resolve-DnsName fail या REFUSED देता है।", "Public DNS 1.1.1.1 से test करें। केवल local DNS fail हो तो adapter DNS 1.1.1.1 और 8.8.8.8 करें, cache clear करें। Domain globally fail हो तो Railway domain status check करें।"),
    ("Blank Shopify iframe", "App title दिखता है, content area blank रहता है।", "App Home URL, callback URL, DNS और HTTP response check करें। Browser hard refresh करें। Railway logs में request आनी चाहिए।"),
    ("OAuth / invalid API key", "Login loop, unauthorized, invalid client या callback mismatch।", "API key + secret को उसी Shopify app से लें। Shopify config और Railway variables को एक ही app identity पर रखें। Secret कभी Git में commit न करें।"),
    ("PostgreSQL error", "Prisma migration या database connection fail।", "DATABASE_URL Railway PostgreSQL internal hostname का हो। Postgres service Online हो। Existing volume को delete न करें।"),
    ("Redis / BullMQ error", "Queue शुरू नहीं होती या ECONNREFUSED 6379।", "REDIS_URL Railway Redis internal hostname/password का हो। Redis service और volume Online हों।"),
]:
    doc.add_paragraph(title, style="Heading 2")
    p=doc.add_paragraph(); set_font(p.add_run("पहचान: "),11,True,RED); set_font(p.add_run(symptoms),11)
    p=doc.add_paragraph(); set_font(p.add_run("Fix: "),11,True,GREEN); set_font(p.add_run(fix),11)

doc.add_page_break()
doc.add_paragraph("4. Config values जो हमेशा match होने चाहिए", style="Heading 1")
table = doc.add_table(rows=1, cols=2); table.style="Table Grid"; set_table_widths(table,[2700,6660])
for c,t in zip(table.rows[0].cells,["Setting","Expected rule"]): set_cell(c,t,True,LIGHT,DARK)
for k,v in [
    ("PORT", "3000; server 0.0.0.0 पर listen करे।"),
    ("SHOPIFY_APP_URL", "Current Railway HTTPS public domain; trailing slash के बिना।"),
    ("Shopify application_url", "SHOPIFY_APP_URL के exact बराबर।"),
    ("OAuth redirect", "<SHOPIFY_APP_URL>/auth/callback"),
    ("SHOPIFY_API_KEY", "‘store duplicate’ Shopify app की Client ID।"),
    ("SHOPIFY_API_SECRET", "उसी ‘store duplicate’ app का secret; chat/Git में कभी न रखें।"),
    ("DATABASE_URL", "Railway Postgres internal URL; app service से accessible।"),
    ("REDIS_URL", "Railway Redis internal URL; app service से accessible।"),
]:
    cells=table.add_row().cells; set_cell(cells[0],k,True); set_cell(cells[1],v)

doc.add_paragraph("क्या कभी delete नहीं करना", style="Heading 2")
for x in [
    "PostgreSQL service या postgres-volume, जब तक verified backup और explicit migration plan न हो।",
    "Redis volume बिना यह समझे कि pending jobs/idempotency state पर क्या प्रभाव पड़ेगा।",
    ".env या secrets को GitHub पर push न करें।",
    "Working domain delete न करें; पहले replacement/rollback plan रखें।",
]: add_bullet(x)

add_callout("Secret leak response:", "यदि कोई Shopify secret, database password या Redis password chat/screenshot में share हो जाए, उसे exposed मानें और provider dashboard से rotate करें।", "FDECEC", RED)

doc.add_paragraph("5. Final health check", style="Heading 1")
for x in [
    "Shopify Admin में store duplicate app dashboard खुलता है।",
    "Railway app deployment SUCCESS है और replica RUNNING है।",
    "Public URL HTTP 200 देता है।",
    "PostgreSQL migration: No pending migrations / no errors।",
    "Redis connection error नहीं है।",
    "Latest logs में repeated crash, 401, 500 या DNS errors नहीं हैं।",
]: add_bullet("☐ " + x)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Last updated: 15 August 2026 • Project: store duplicate"),9,color="6B7280")

doc.core_properties.title = "Store Duplicate Troubleshooting & Recovery Guide"
doc.core_properties.subject = "Shopify and Railway recovery runbook"
doc.core_properties.author = "Store Duplicate Operations"
doc.save(OUT)
print(str(OUT))
